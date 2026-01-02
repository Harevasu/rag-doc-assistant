"""
Document Processor Module.
Handles extraction of text from PDFs, Word docs, and text files, and chunking documents.
"""

import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from typing import List, Dict, Any


from config import CHUNK_SIZE, CHUNK_OVERLAP


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as a string
    """
    text_content = []
    
    with fitz.open(file_path) as doc:
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            if page_text.strip():
                text_content.append(f"[Page {page_num}]\n{page_text}")
    
    return "\n\n".join(text_content)


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text content from a text file.
    
    Args:
        file_path: Path to the text file
        
    Returns:
        File content as a string
    """
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a Word document (.docx).
    
    Args:
        file_path: Path to the Word document
        
    Returns:
        Extracted text as a string
    """
    doc = Document(file_path)
    text_content = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_content.append(" | ".join(row_text))
    
    return "\n\n".join(text_content)


def extract_text(file_path: str) -> str:
    """
    Extract text from a file based on its extension.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file type is not supported
    """
    path = Path(file_path)
    extension = path.suffix.lower()
    
    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif extension == ".docx":
        return extract_text_from_docx(file_path)
    elif extension in [".txt", ".md", ".text"]:
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}")


def chunk_document(
    text: str,
    document_id: str,
    filename: str,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP
) -> List[Dict[str, Any]]:
    """
    Split document text into overlapping chunks with metadata.
    
    Args:
        text: The full document text
        document_id: Unique identifier for the document
        filename: Original filename
        chunk_size: Maximum size of each chunk
        chunk_overlap: Number of characters to overlap between chunks
        
    Returns:
        List of chunk dictionaries with text and metadata
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    chunks = text_splitter.split_text(text)
    
    chunk_documents = []
    for i, chunk_text in enumerate(chunks):
        chunk_documents.append({
            "id": f"{document_id}_chunk_{i}",
            "text": chunk_text,
            "metadata": {
                "document_id": document_id,
                "filename": filename,
                "chunk_index": i,
                "total_chunks": len(chunks)
            }
        })
    
    return chunk_documents


def process_document(file_path: str, document_id: str) -> List[Dict[str, Any]]:
    """
    Full pipeline to process a document: extract text and chunk it.
    
    Args:
        file_path: Path to the document file
        document_id: Unique identifier for the document
        
    Returns:
        List of processed chunks with metadata
    """
    filename = Path(file_path).name
    text = extract_text(file_path)
    chunks = chunk_document(text, document_id, filename)
    
    return chunks
